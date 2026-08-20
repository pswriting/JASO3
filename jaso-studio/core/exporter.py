# -*- coding: utf-8 -*-
"""자소서 스튜디오 — DOCX / TXT 내보내기"""
import io
import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
GOLD = RGBColor(0x8A, 0x6D, 0x3F)
INK = RGBColor(0x20, 0x24, 0x2B)
GRAY = RGBColor(0x6E, 0x68, 0x57)


def _set_korean_font(run, name="맑은 고딕", size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)


def build_docx(company: str, role: str, items: list) -> bytes:
    """
    items: [{question, subtitle, body, chars_incl, chars_excl, limit, count_mode}]
    """
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

    # 표지 헤더
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("자기소개서")
    _set_korean_font(r, size=20, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_parts = [x for x in [company.strip(), role.strip()] if x]
    r = p.add_run(" · ".join(meta_parts) if meta_parts else "")
    _set_korean_font(r, size=11, color=GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(datetime.date.today().strftime("%Y. %m. %d"))
    _set_korean_font(r, size=9, color=GRAY)

    doc.add_paragraph()

    for i, item in enumerate(items, 1):
        # 문항
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{i}. {item['question']}")
        _set_korean_font(r, size=11.5, bold=True, color=NAVY)

        # 소제목
        if item.get("subtitle"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(f"[{item['subtitle']}]")
            _set_korean_font(r, size=11, bold=True, color=GOLD)

        # 본문
        for para in str(item.get("body", "")).split("\n"):
            para = para.strip()
            if not para:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.55
            p.paragraph_format.space_after = Pt(7)
            r = p.add_run(para)
            _set_korean_font(r, size=10.5, color=INK)

        # 글자수 캡션
        mode_txt = "공백 포함" if item.get("count_mode") == "incl" else "공백 제외"
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(
            f"— {mode_txt} {item.get('chars_incl' if item.get('count_mode') == 'incl' else 'chars_excl', 0)}자"
            f" / 제한 {item.get('limit', '?')}자"
            f"  (공백 포함 {item.get('chars_incl', 0)} · 공백 제외 {item.get('chars_excl', 0)})")
        _set_korean_font(r, size=8.5, color=GRAY)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_txt(company: str, role: str, items: list) -> str:
    lines = []
    header = " · ".join(x for x in [company.strip(), role.strip()] if x)
    if header:
        lines += [header, "=" * 40, ""]
    for i, item in enumerate(items, 1):
        lines.append(f"{i}. {item['question']}")
        if item.get("subtitle"):
            lines.append(f"[{item['subtitle']}]")
        lines.append(str(item.get("body", "")).strip())
        mode_txt = "공백 포함" if item.get("count_mode") == "incl" else "공백 제외"
        n = item.get("chars_incl" if item.get("count_mode") == "incl" else "chars_excl", 0)
        lines.append(f"({mode_txt} {n}자 / 제한 {item.get('limit', '?')}자)")
        lines.append("")
    return "\n".join(lines)
