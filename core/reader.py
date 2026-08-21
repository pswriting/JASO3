# -*- coding: utf-8 -*-
"""업로드 파일(이력서·경력기술서·기존 자소서) 텍스트 자동 추출"""
import io
import re
import zipfile

SUPPORTED = ["pdf", "docx", "txt", "md", "hwp", "hwpx"]


def extract_text(filename: str, data: bytes):
    """반환: (추출 텍스트, 경고 메시지)"""
    name = (filename or "").lower()
    try:
        if name.endswith(".docx"):
            return _from_docx(data), ""
        if name.endswith(".pdf"):
            return _from_pdf(data)
        if name.endswith((".txt", ".md")):
            return data.decode("utf-8", errors="replace").strip(), ""
        if name.endswith(".hwpx"):
            return _from_hwpx(data)
        if name.endswith(".hwp"):
            return _from_hwp(data)
        return "", "지원하지 않는 형식입니다 (pdf·docx·txt·hwp·hwpx 지원)"
    except Exception as e:
        return "", f"추출 실패: {e}"


def _from_docx(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(dict.fromkeys(cells)))
    return "\n".join(parts).strip()


def _from_pdf(data: bytes):
    import pdfplumber
    out, warn = [], ""
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        pages = pdf.pages[:20]
        if len(pdf.pages) > 20:
            warn = "20페이지까지만 추출했습니다"
        for page in pages:
            out.append(page.extract_text() or "")
    text = "\n".join(out).strip()
    if not text:
        warn = "텍스트를 찾지 못했습니다 (스캔 이미지 PDF일 수 있어요)"
    return text, warn


def _from_hwpx(data: bytes):
    zf = zipfile.ZipFile(io.BytesIO(data))
    texts = []
    for n in sorted(zf.namelist()):
        if n.startswith("Contents/section") and n.endswith(".xml"):
            xml = zf.read(n).decode("utf-8", errors="replace")
            xml = re.sub(r"<hp:t[^>]*>", "", xml)
            xml = re.sub(r"<[^>]+>", "", xml)
            texts.append(xml.replace("", " "))
    text = re.sub(r"[ \t]+", " ", "\n".join(texts)).strip()
    return text, ("" if text else "본문을 찾지 못했습니다")


def _from_hwp(data: bytes):
    import olefile
    ole = olefile.OleFileIO(io.BytesIO(data))
    try:
        if ole.exists("PrvText"):
            raw = ole.openstream("PrvText").read()
            text = raw.decode("utf-16-le", errors="replace").strip("\x00 \r\n")
            if text.strip():
                return text.strip(), "HWP는 미리보기 기반이라 일부만 추출될 수 있어요 — 가능하면 PDF로 변환해 올려 주세요"
        return "", "이 HWP에서 텍스트를 추출하지 못했습니다 — PDF로 변환해 올려 주세요"
    finally:
        ole.close()
