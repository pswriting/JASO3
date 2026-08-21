# -*- coding: utf-8 -*-
"""합격 이력서 3종 + 경력기술서 양식 (DOCX — 한글(HWP)에서도 열림)"""
import io

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
GOLD = RGBColor(0x8A, 0x6D, 0x3F)
INK = RGBColor(0x20, 0x24, 0x2B)
GRAY = RGBColor(0x8A, 0x85, 0x76)
HINT = RGBColor(0xA8, 0xA2, 0x92)

NAVY_HEX, GOLD_HEX, LINE_HEX, IVORY_HEX = "1B2A4A", "8A6D3F", "DDD5C4", "F5F1E6"


# ── 공통 헬퍼 ──────────────────────────────

def _font(run, size=10, bold=False, color=INK, name="맑은 고딕"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)


def _para(doc, text="", size=10, bold=False, color=INK, align=None,
          before=0, after=4, name="맑은 고딕"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    if text:
        _font(p.add_run(text), size=size, bold=bold, color=color, name=name)
    return p


def _rule(doc, color=GOLD_HEX, size="12", after=8):
    """단락 하단 경계선(수평선)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    ppr = p._element.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    ppr.append(borders)
    return p


def _section_title(doc, text):
    p = _para(doc, text, size=12, bold=True, color=NAVY, before=14, after=2)
    ppr = p._element.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), LINE_HEX)
    borders.append(bottom)
    ppr.append(borders)


def _shade(cell, hex_color):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcpr.append(shd)


def _table_borders(table, color=LINE_HEX):
    tbl = table._tbl
    tblpr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblpr.append(borders)


def _cell_text(cell, text, size=9.5, bold=False, color=INK, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if align is not None:
        p.alignment = align
    _font(p.add_run(text), size=size, bold=bold, color=color)


def _header_row(table, labels):
    for i, label in enumerate(labels):
        cell = table.rows[0].cells[i]
        _shade(cell, NAVY_HEX)
        _cell_text(cell, label, size=9.5, bold=True,
                   color=RGBColor(0xF5, 0xF1, 0xE6), align=WD_ALIGN_PARAGRAPH.CENTER)


def _hint_rows(table, rows):
    """rows: [[텍스트,...], ...] — 연회색 힌트"""
    for r, row_vals in enumerate(rows, start=1):
        for c, val in enumerate(row_vals):
            _cell_text(table.rows[r].cells[c], val, size=9.5, color=HINT)


def _new_doc():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.8)
        s.bottom_margin = Cm(1.8)
        s.left_margin = Cm(2.0)
        s.right_margin = Cm(2.0)
    return doc


def _tip(doc, text):
    _para(doc, "✎ " + text, size=8.5, color=GRAY, before=6, after=2)


def _to_bytes(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── ① 클래식 정갈형 (대기업·공기업) ──────────

def build_resume_classic() -> bytes:
    doc = _new_doc()
    _para(doc, "이  력  서", size=22, bold=True, color=NAVY,
          align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    _para(doc, "CLASSIC RESUME — 대기업·공기업 서류 표준형", size=8.5, color=GOLD,
          align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    _rule(doc)

    _section_title(doc, "인적 사항")
    t = doc.add_table(rows=4, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _table_borders(t)
    labels = [("성명", "홍길동", "생년", "1995년 (만 30세)"),
              ("연락처", "010-0000-0000", "이메일", "example@email.com"),
              ("주소", "서울특별시 ○○구 (통근 30분권이면 구 단위까지만)", "", ""),
              ("지원 부문", "○○기업 ○○직무 신입/경력", "", "")]
    for r, (l1, v1, l2, v2) in enumerate(labels):
        _shade(t.rows[r].cells[0], IVORY_HEX)
        _cell_text(t.rows[r].cells[0], l1, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(t.rows[r].cells[1], v1, color=HINT)
        if l2:
            _shade(t.rows[r].cells[2], IVORY_HEX)
            _cell_text(t.rows[r].cells[2], l2, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell_text(t.rows[r].cells[3], v2, color=HINT)
        else:
            t.rows[r].cells[1].merge(t.rows[r].cells[3])

    _section_title(doc, "학력 사항")
    t = doc.add_table(rows=3, cols=4)
    _table_borders(t)
    _header_row(t, ["기간", "학교명", "전공", "비고(학점/졸업구분)"])
    _hint_rows(t, [["2015.03 ~ 2021.02", "○○대학교", "경영학과", "3.8 / 4.5 · 졸업"],
                   ["2012.03 ~ 2015.02", "○○고등학교", "-", "졸업"]])
    _tip(doc, "최종 학력부터 역순으로. 학점은 3.5 이상일 때만 기재하는 것이 유리합니다.")

    _section_title(doc, "경력 사항")
    t = doc.add_table(rows=3, cols=4)
    _table_borders(t)
    _header_row(t, ["기간", "회사명", "직책/부서", "핵심 성과 한 줄"])
    _hint_rows(t, [["2023.01 ~ 재직", "○○주식회사", "영업팀 주임", "담당 채널 매출 전년 대비 32% 증대"],
                   ["2021.03 ~ 2022.12", "○○기업", "마케팅팀 사원", "광고 ROAS 1,600% 달성"]])
    _tip(doc, "성과는 반드시 숫자로. '무엇을 했다'가 아니라 '무엇이 좋아졌다'를 쓰세요.")

    _section_title(doc, "자격·어학")
    t = doc.add_table(rows=3, cols=3)
    _table_borders(t)
    _header_row(t, ["취득일", "자격/시험명", "점수·등급"])
    _hint_rows(t, [["2024.05", "OPIc", "IH"], ["2023.11", "컴퓨터활용능력", "1급"]])

    _section_title(doc, "수상·대외활동")
    t = doc.add_table(rows=2, cols=3)
    _table_borders(t)
    _header_row(t, ["시기", "내용", "주최/규모"])
    _hint_rows(t, [["2024.09", "○○ 공모전 대상", "○○부 주최 · 참가 300팀"]])

    _para(doc, "위 내용은 사실과 다름이 없습니다.", size=9, color=INK,
          align=WD_ALIGN_PARAGRAPH.CENTER, before=18)
    _para(doc, "2026년   월   일        지원자  홍길동  (인)", size=9, color=INK,
          align=WD_ALIGN_PARAGRAPH.CENTER)
    return _to_bytes(doc)


# ── ② 모던 미니멀형 (IT·스타트업·중고신입) ────

def build_resume_modern() -> bytes:
    doc = _new_doc()
    _para(doc, "홍길동", size=24, bold=True, color=NAVY, after=0)
    _para(doc, "데이터로 매출을 만드는 퍼포먼스 마케터  ←  나를 한 문장으로 정의하세요",
          size=10.5, color=GOLD, after=2)
    _para(doc, "010-0000-0000  ·  example@email.com  ·  서울  ·  포트폴리오 링크",
          size=9, color=GRAY, after=0)
    _rule(doc)

    _section_title(doc, "핵심 역량")
    for line in ["▸  퍼포먼스 광고 운영 3년 — 누적 집행 12억, 평균 ROAS 480%",
                 "▸  SQL·GA4 기반 고객 데이터 분석 — 이탈 구간 발견해 전환율 1.8배",
                 "▸  콘텐츠 기획 — 오가닉 유입 월 2만 → 11만 성장"]:
        _para(doc, line, size=10, color=HINT, after=3)
    _tip(doc, "지원 직무의 JD 키워드와 정확히 겹치는 역량 3개만. 각 줄 끝은 반드시 숫자로.")

    _section_title(doc, "경력")
    _para(doc, "○○컴퍼니  ·  그로스마케팅팀 매니저", size=11, bold=True, color=INK, before=4, after=1)
    _para(doc, "2023.01 — 재직 중", size=8.5, color=GRAY, after=3)
    for line in ["–  신규 채널 발굴로 획득 단가(CAC) 34% 절감",
                 "–  A/B 테스트 47회 운영, 랜딩 전환율 2.1% → 3.9%"]:
        _para(doc, line, size=9.5, color=HINT, after=2)
    _para(doc, "○○스타트업  ·  마케팅 인턴", size=11, bold=True, color=INK, before=8, after=1)
    _para(doc, "2022.01 — 2022.12", size=8.5, color=GRAY, after=3)
    _para(doc, "–  SNS 콘텐츠 90건 제작, 팔로워 0 → 1.2만", size=9.5, color=HINT, after=2)

    _section_title(doc, "학력")
    _para(doc, "○○대학교  경영학과  ·  2015.03 — 2021.02", size=9.5, color=HINT, after=2)

    _section_title(doc, "기술·자격")
    _para(doc, "GA4  ·  SQL(중급)  ·  Meta/Google Ads  ·  Figma  ·  OPIc IH", size=9.5, color=HINT, after=2)
    _tip(doc, "이 양식은 사진·주소 없이 역량으로 승부하는 형식입니다. IT·스타트업·수시채용에 적합합니다.")
    return _to_bytes(doc)


# ── ③ 성과 강조형 (경력직 이직) ───────────────

def build_resume_performance() -> bytes:
    doc = _new_doc()
    t = doc.add_table(rows=1, cols=1)
    _table_borders(t, color=NAVY_HEX)
    cell = t.rows[0].cells[0]
    _shade(cell, NAVY_HEX)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    _font(p.add_run("홍길동  |  영업관리  ·  경력 5년"), size=15, bold=True,
          color=RGBColor(0xF5, 0xF1, 0xE6))
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(8)
    _font(p2.add_run("010-0000-0000  ·  example@email.com"), size=9,
          color=RGBColor(0xC9, 0xC2, 0xAE))

    _section_title(doc, "커리어 요약")
    _para(doc, "호텔·유통 세일즈 5년. 데이터 기반 고객 분석으로 3개 회사에서 모두 담당 매출 신기록을 세웠습니다. "
               "숫자로 검증된 영업 기획자입니다.  ← 3줄 이내, 숫자 포함, 두괄식으로.",
          size=10, color=HINT, after=4)

    _section_title(doc, "대표 성과  TOP 3")
    for line in ["①  VIP 전시 행사 단독 기획 — 행사 매출 80억 원 (전사 분기 최고 기록)",
                 "②  신규 멤버십 설계로 월 매출 3,500만 원 → 8,000만 원 (128%↑)",
                 "③  고객 데이터 세분화로 재구매율 12% → 31%"]:
        _para(doc, line, size=10.5, bold=True, color=INK, after=3)
    _tip(doc, "인사담당자는 여기까지만 읽고 서류를 결정합니다. 가장 큰 숫자 3개를 올리세요.")

    _section_title(doc, "경력 상세")
    t = doc.add_table(rows=3, cols=4)
    _table_borders(t)
    _header_row(t, ["기간", "회사/직책", "담당 업무", "핵심 성과(수치)"])
    _hint_rows(t, [["2023.01 ~ 재직", "○○리테일 · 영업기획 대리", "VIP 고객 관리, 행사 기획", "행사 매출 80억, 신규 VIP 20%↑"],
                   ["2020.02 ~ 2022.12", "○○호텔 · 멤버십 세일즈", "멤버십 기획·판매", "월 매출 128%↑, 최우수 직원"]])

    _section_title(doc, "학력·자격")
    _para(doc, "○○대학교 경영학과 졸업 (2020)  ·  유통관리사 2급  ·  OPIc IH", size=9.5, color=HINT, after=2)
    _tip(doc, "경력직은 '무엇을 담당했는가'보다 '무엇을 바꿔놓았는가'가 전부입니다.")
    return _to_bytes(doc)


# ── 경력기술서 ────────────────────────────────

def build_career_sheet() -> bytes:
    doc = _new_doc()
    _para(doc, "경 력 기 술 서", size=20, bold=True, color=NAVY,
          align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    _para(doc, "CAREER DESCRIPTION — 회사별 · 성과 중심", size=8.5, color=GOLD,
          align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    _rule(doc)
    _para(doc, "작성 원칙  ①역할이 아니라 '내가 한 행동'을 쓴다  ②모든 성과는 숫자로 끝낸다  "
               "③지원 직무와 무관한 업무는 과감히 뺀다", size=9, color=GRAY, after=10)

    for idx in ("1", "2"):
        _section_title(doc, f"경력 {idx}")
        t = doc.add_table(rows=3, cols=4)
        _table_borders(t)
        for r, (l1, v1, l2, v2) in enumerate([
                ("회사명", "○○주식회사 (업종: 유통)", "재직 기간", "2023.01 ~ 재직 (2년 8개월)"),
                ("소속/직책", "영업기획팀 · 대리", "최종 연봉/직급", "선택 기재"),
                ("퇴직 사유", "이직 준비 중 (경력 1은 재직 중이면 생략)", "", "")]):
            _shade(t.rows[r].cells[0], IVORY_HEX)
            _cell_text(t.rows[r].cells[0], l1, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell_text(t.rows[r].cells[1], v1, color=HINT)
            if l2:
                _shade(t.rows[r].cells[2], IVORY_HEX)
                _cell_text(t.rows[r].cells[2], l2, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
                _cell_text(t.rows[r].cells[3], v2, color=HINT)
            else:
                t.rows[r].cells[1].merge(t.rows[r].cells[3])

        _para(doc, "담당 업무", size=10.5, bold=True, color=NAVY, before=8, after=2)
        for line in ["–  VIP 고객 데이터 분석 및 등급별 관리 전략 수립",
                     "–  분기 판촉 행사 기획·운영 (연 8회, 회당 예산 2억)",
                     "–  매장 20개점 영업 실적 관리 및 개선 코칭"]:
            _para(doc, line, size=9.5, color=HINT, after=2)

        _para(doc, "주요 성과", size=10.5, bold=True, color=NAVY, before=6, after=2)
        for line in ["–  VIP 전시 행사 단독 기획 → 행사 매출 80억 원, 전사 분기 최고 기록",
                     "–  고객 세분화 후 맞춤 제안 도입 → 재구매율 12% → 31%",
                     "–  신규 멤버십 설계 → 월 매출 3,500만 원 → 8,000만 원"]:
            _para(doc, line, size=9.5, color=HINT, after=2)

        _para(doc, "활용 역량·도구", size=10.5, bold=True, color=NAVY, before=6, after=2)
        _para(doc, "Excel(피벗·대시보드) · Salesforce · SQL 기초 · PB상품 소싱 협상", size=9.5, color=HINT, after=8)

    _tip(doc, "회사가 많으면 지원 직무와 관련 깊은 경력 2~3개만 상세히 쓰고, 나머지는 한 줄로 요약하세요.")
    _tip(doc, "이 파일은 워드(.docx)이며 한글(HWP)에서도 그대로 열립니다. 회색 예시를 본인 내용으로 바꿔 쓰세요.")
    return _to_bytes(doc)


TEMPLATES = {
    "classic": ("이력서 ① 클래식 정갈형", "대기업·공기업 표준 — 표 기반 정석 구성", build_resume_classic),
    "modern": ("이력서 ② 모던 미니멀형", "IT·스타트업·수시채용 — 역량 중심, 사진·주소 없음", build_resume_modern),
    "performance": ("이력서 ③ 성과 강조형", "경력직 이직 — 대표 성과 TOP 3가 먼저 보이는 구성", build_resume_performance),
    "career": ("경력기술서", "회사별 · 성과 중심 — 담당 업무/성과/도구 구조", build_career_sheet),
}
