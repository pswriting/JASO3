# -*- coding: utf-8 -*-
"""자소서 스튜디오 — DOCX / TXT 내보내기 (프리미엄 문서 디자인)"""
import io
import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
GOLD = RGBColor(0x8A, 0x6D, 0x3F)
CHAMPAGNE = RGBColor(0xC9, 0xC2, 0xAE)
IVORY = RGBColor(0xF5, 0xF1, 0xE6)
INK = RGBColor(0x22, 0x26, 0x2D)
GRAY = RGBColor(0x8A, 0x85, 0x76)

NAVY_HEX, GOLD_HEX, LINE_HEX = "1B2A4A", "8A6D3F", "DDD5C4"


def _font(run, size=10.5, bold=False, color=INK, name="맑은 고딕"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)


def _bottom_border(p, color=GOLD_HEX, sz="8"):
    ppr = p._element.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    ppr.append(borders)


def _shade_cell(cell, hex_color):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcpr.append(shd)


def _band_borders(table, color=NAVY_HEX):
    tblpr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblpr.append(borders)


def _page_number_footer(doc):
    """가운데 정렬 페이지 번호."""
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    _font(run, size=8, color=GRAY)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_end)


def build_docx(company: str, role: str, items: list) -> bytes:
    """
    items: [{question, subtitle, body, chars_incl, chars_excl, limit, count_mode}]
    """
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.3)
        section.right_margin = Cm(2.3)

    # ── 표지 헤더 밴드 (네이비) ──
    band = doc.add_table(rows=1, cols=1)
    _band_borders(band)
    cell = band.rows[0].cells[0]
    _shade_cell(cell, NAVY_HEX)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(2)
    _font(p.add_run("자  기  소  개  서"), size=19, bold=True, color=IVORY, name="맑은 고딕")

    meta_parts = [x for x in [(company or "").strip(), (role or "").strip()] if x]
    if meta_parts:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(2)
        _font(p2.add_run("  ·  ".join(meta_parts)), size=10.5, color=CHAMPAGNE)
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(14)
    _font(p3.add_run(datetime.date.today().strftime("%Y. %m. %d")), size=8, color=CHAMPAGNE)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)

    # ── 문항 블록 ──
    for i, item in enumerate(items, 1):
        # QUESTION 번호 (골드 오버라인)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16 if i > 1 else 4)
        p.paragraph_format.space_after = Pt(1)
        _font(p.add_run(f"Q U E S T I O N   {i:02d}"), size=8.5, bold=True, color=GOLD)

        # 문항 전문 + 골드 밑줄
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        _bottom_border(p, color=GOLD_HEX, sz="6")
        _font(p.add_run(item["question"]), size=12, bold=True, color=NAVY)

        # 소제목 — 골드 대괄호 + 네이비 제목
        if item.get("subtitle"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            _font(p.add_run("["), size=12.5, bold=True, color=GOLD)
            _font(p.add_run(item["subtitle"]), size=12.5, bold=True, color=NAVY)
            _font(p.add_run("]"), size=12.5, bold=True, color=GOLD)

        # 본문 — 양쪽 정렬, 행간 1.6
        for para in str(item.get("body", "")).split("\n"):
            para = para.strip()
            if not para:
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.6
            p.paragraph_format.space_after = Pt(9)
            _font(p.add_run(para), size=10.5, color=INK)

        # 글자수 캡션 — 오른쪽 정렬
        mode_txt = "공백 포함" if item.get("count_mode") == "incl" else "공백 제외"
        n = item.get("chars_incl" if item.get("count_mode") == "incl" else "chars_excl", 0)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(2)
        _font(p.add_run(f"{mode_txt} {n:,}자 / 제한 {item.get('limit', '?')}자"),
              size=8, color=GRAY)

    _page_number_footer(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_txt(company: str, role: str, items: list) -> str:
    lines = []
    header = " · ".join(x for x in [company.strip(), role.strip()] if x)
    if header:
        lines += [header, "=" * 40, ""]
    for i, item in enumerate(items, 1):
        lines.append(f"{i}. {item['question']}")
        if item.get("subtitle"):
            lines.append(f"[{item['subtitle']}]")
        lines.append(str(item.get("body", "")).strip())
        mode_txt = "공백 포함" if item.get("count_mode") == "incl" else "공백 제외"
        n = item.get("chars_incl" if item.get("count_mode") == "incl" else "chars_excl", 0)
        lines.append(f"({mode_txt} {n}자 / 제한 {item.get('limit', '?')}자)")
        lines.append("")
    return "\n".join(lines)
